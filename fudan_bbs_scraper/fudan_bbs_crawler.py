import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
from datetime import datetime
import random
import re

# 基本配置
BASE_XML_URL = "https://bbs.fudan.edu.cn/v18/0an?path=/groups/sport.faq/Running/DACA5FD39/DA14D15FC/D5B8DAE31/DA78613A9"
BASE_ANC_PATH = "/groups/sport.faq/Running/DACA5FD39/DA14D15FC/D5B8DAE31/DA78613A9"
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

def extract_bbs_content(xml_url):
    """从复旦大学BBS XML页面提取帖子列表（标题、作者、时间、链接）"""
    try:
        response = requests.get(xml_url, headers=HEADERS, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'xml')
        
        posts = []
        for ent in soup.find_all('ent'):
            title = ent.get_text(strip=True)
            path = ent.get('path', '')
            author = ent.get('id', '')
            post_time = ent.get('time', '')
            
            if title and path and path.startswith('/M.'):
                full_url = f"https://bbs.fudan.edu.cn/v18/anc?path={BASE_ANC_PATH}{path}"
                
                # 格式化帖子信息
                post_info = f"{title}"
                info_parts = []
                if author:
                    info_parts.append(f"作者: {author}")
                if post_time:
                    try:
                        dt = datetime.fromisoformat(post_time.replace('Z', '+00:00'))
                        formatted_time = dt.strftime('%Y-%m-%d %H:%M:%S')
                        info_parts.append(f"时间: {formatted_time}")
                    except:
                        info_parts.append(f"时间: {post_time}")
                if info_parts:
                    post_info += f" ({', '.join(info_parts)})"
                
                posts.append((post_info, full_url))
        return posts
    
    except Exception as e:
        print(f"抓取 XML 页面出错: {e}")
        return []

def clean_post_content(text):
    """清理帖子正文，去掉底部签名、广告和多余空行"""
    # 去掉连续多行空白
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    
    # 去掉常见签名/广告，例如 "-- " 开头的行或末尾"发自"之类的行
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line_strip = line.strip()
        if (line_strip.startswith('--') or 
            line_strip.startswith('※ 来源:') or
            '发自' in line_strip or 
            '来自' in line_strip or
            line_strip.startswith('※ 修改:')):
            continue
        cleaned_lines.append(line)
    
    cleaned_text = '\n'.join(cleaned_lines).strip()
    return cleaned_text

def extract_post_text(url):
    """抓取帖子正文，针对复旦大学BBS特定页面结构"""
    try:
        response = requests.get(url, headers=HEADERS, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # 尝试多种可能的正文容器
        content_selectors = [
            'pre',  # 复旦大学BBS常用格式
            'article',
            'div.content',
            'div.post-content',
            'div.article-content'
        ]
        
        content = None
        for selector in content_selectors:
            element = soup.select_one(selector)
            if element:
                content = element.get_text(separator='\n', strip=True)
                break
        
        # 如果上述选择器都没找到，尝试查找包含大量文本的容器
        if not content:
            # 查找所有可能包含正文的div
            potential_divs = soup.find_all('div', class_=lambda x: x != 'header' and x != 'footer')
            for div in potential_divs:
                text = div.get_text(separator='\n', strip=True)
                # 如果文本长度较长，可能是正文
                if len(text) > 30:
                    content = text
                    break
        
        if content:
            content = clean_post_content(content)
            print(f"[{url}] 正文抓取成功")
            return content
        else:
            print(f"[{url}] 无法获取正文内容，尝试查看页面结构...")
            # 打印页面标题和前500字符以便调试
            title = soup.find('title')
            if title:
                print(f"页面标题: {title.get_text()}")
            print(f"页面预览: {str(soup)[:5000]}...")
            return "[无法获取正文内容]"
            
    except Exception as e:
        print(f"[{url}] 抓取正文失败, 错误: {e}")
        return f"[抓取正文失败: {str(e)}]"

def create_word_document(posts, filename):
    """创建Word文档，将标题、作者、时间和正文内容写入"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('复旦大学BBS跑步版块内容提取', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 说明段落
    doc.add_paragraph(
        f"此文档包含从复旦大学BBS跑步版块提取的所有帖子内容，共 {len(posts)} 个条目。"
        f"\n提取时间: {time.strftime('%Y-%m-%d %H:%M:%S')}"
    )
    doc.add_paragraph()
    
    total = len(posts)
    for i, (post_info, url) in enumerate(posts, 1):
        print(f"\n正在抓取帖子 {i}/{total} ...")
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(post_info)
        
        # 添加URL
        url_para = doc.add_paragraph()
        url_para.add_run("链接: ").bold = True
        url_para.add_run(url)
        
        # 抓取正文
        content = extract_post_text(url)
        content_para = doc.add_paragraph(content)
        content_para.paragraph_format.space_after = Pt(6)
        
        # 分隔线
        doc.add_paragraph("─" * 80)
        
        # 随机延时，防止访问太快
        time.sleep(random.uniform(0.1, 0.5))
    
    doc.save(filename)
    print(f"\n文档已保存为: {filename}")

def main():
    print("开始提取复旦大学BBS内容...")
    posts = extract_bbs_content(BASE_XML_URL)
    
    if posts:
        print(f"成功提取 {len(posts)} 个帖子列表")
        # 只处理前几个帖子作为测试
        # posts = posts[:5]  # 取消注释这行来测试少量帖子
        filename = "复旦大学BBS跑步版块内容.docx"
        create_word_document(posts, filename)
        print("\n提取完成！")
        print(f"共提取 {len(posts)} 个帖子")
        print(f"文档已保存为: {filename}")
    else:
        print("未能提取到任何内容")

if __name__ == "__main__":
    main()